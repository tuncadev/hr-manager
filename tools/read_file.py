import io

import PyPDF2
from docx import Document
import chardet


class ReadFileContents:
    def __init__(self, file_path, file_content=None):
        self.file_path = file_path
        self.file_content = file_content

    def read_doc_file(self):
        if self.file_content:
            doc = Document(io.BytesIO(self.file_content))
        else:
            doc = Document(self.file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return text

    def read_docx_file(self):
        if self.file_content:
            doc = Document(io.BytesIO(self.file_content))
        else:
            doc = Document(self.file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return text

    def read_txt_file(self):
        text = ""
        if self.file_content:
            try:
                text = self.file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text = self.file_content.decode('utf-16le')
                except UnicodeDecodeError:
                    try:
                        text = self.file_content.decode('utf-32')
                    except UnicodeDecodeError:
                        return "Error: File content is not readable."
        else:
            with open(self.file_path, 'rb') as file:
                raw_data = file.read(1000)
                encoding_info = chardet.detect(raw_data)
                encoding = encoding_info.get('encoding')
                if encoding:
                    try:
                        text = file.read().decode(encoding)
                    except UnicodeDecodeError:
                        try:
                            text = file.read().decode('utf-16le')
                        except UnicodeDecodeError:
                            try:
                                text = file.read().decode('utf-32')
                            except UnicodeDecodeError:
                                return "Error: File content is not readable."
        return text

   # def read_pdf_file(self):
    #    return extract_text(io.BytesIO(self.file_content))

    def read_pdf_file(self, file_content):
        # Since 'file_content' is already a byte stream from Streamlit
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

        # Extract text from all pages
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()

        return text

    def check_for_binary_data(self):
        if self.file_content:
            try:
                self.file_content.decode('utf-8')
            except UnicodeDecodeError:
                return "Error: File contains binary data or is not readable."
        return None